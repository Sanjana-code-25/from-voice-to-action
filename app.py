import io
import json
import os
import tempfile
import urllib.error
import urllib.request
from urllib.parse import quote
from datetime import datetime, timedelta
from typing import Any

import streamlit as st
from dotenv import load_dotenv

load_dotenv()


DEMO_TRANSCRIPT = (
    "We need to tighten the launch plan before Friday. Priya will confirm the final pricing "
    "with finance tomorrow. I will rewrite the onboarding email and send it to the team by "
    "Thursday afternoon. Let's also schedule a 30-minute review with design next week."
)
HISTORY_FILE = os.path.join(os.path.dirname(__file__), "notes_history.json")
MESSAGES_FILE = os.path.join(os.path.dirname(__file__), "teacher_messages.json")
USERS = {
    "teacher1": {"password": "teacher123", "role": "Teacher"},
    "student1": {"password": "student123", "role": "Student"},
}


def load_history() -> list[dict[str, Any]]:
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as history_file:
            return json.load(history_file)
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def save_history(history: list[dict[str, Any]]) -> None:
    with open(HISTORY_FILE, "w", encoding="utf-8") as history_file:
        json.dump(history, history_file, ensure_ascii=False, indent=2)


def load_messages() -> list[dict[str, Any]]:
    try:
        with open(MESSAGES_FILE, "r", encoding="utf-8") as messages_file:
            return json.load(messages_file)
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def save_messages(messages: list[dict[str, Any]]) -> None:
    with open(MESSAGES_FILE, "w", encoding="utf-8") as messages_file:
        json.dump(messages, messages_file, ensure_ascii=False, indent=2)


def server_api_key(name: str) -> str:
    try:
        return str(st.secrets.get(name, "")) or os.getenv(name, "")
    except (FileNotFoundError, KeyError):
        return os.getenv(name, "")


def demo_result(file_name: str) -> dict[str, Any]:
    return {
        "file_name": file_name,
        "transcript": DEMO_TRANSCRIPT,
        "summary": "The launch plan needs a final pricing check, an updated onboarding email, and a design review before the next release milestone.",
        "action_items": [
            {"task": "Confirm final pricing with finance", "owner": "Priya", "due": "Tomorrow"},
            {"task": "Rewrite and send the onboarding email", "owner": "You", "due": "Thursday afternoon"},
            {"task": "Schedule a 30-minute review with design", "owner": "Unassigned", "due": "Next week"},
        ],
    }


def parse_llm_result(content: str, file_name: str, transcript: str) -> dict[str, Any]:
    parsed = json.loads(content)
    return {
        "file_name": file_name,
        "transcript": transcript,
        "summary": parsed.get("summary", "No summary was returned."),
        "action_items": parsed.get("action_items", []),
    }


def process_audio_gemini(uploaded_file: Any, teacher_focus: str, target_language: str) -> dict[str, Any]:
    api_key = server_api_key("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("The app administrator has not configured GEMINI_API_KEY.")

    from google import genai
    from google.genai import types

    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1] or ".wav") as audio_file:
        audio_file.write(uploaded_file.getvalue())
        audio_path = audio_file.name

    try:
        client = genai.Client(api_key=api_key)
        uploaded_audio = client.files.upload(file=audio_path)
        prompt = (
            "Listen to this teacher voice note and return JSON with exactly these keys: "
            "transcript (string in the original spoken language), summary (string), "
            "action_items (array of objects with task, owner, due). "
            "Use 'Unassigned' or 'Not specified' when details are missing. "
            f"Focus on {teacher_focus.lower()} and write summary/action items in {target_language}. "
            "Do not invent names, dates, or tasks."
        )
        response = client.models.generate_content(
            model=os.getenv("GEMINI_MODEL", "gemini-2.0-flash"),
            contents=[uploaded_audio, prompt],
            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.2),
        )
        parsed = json.loads(response.text or "{}")
        return parse_llm_result(
            json.dumps(parsed),
            uploaded_file.name,
            parsed.get("transcript", "No transcript was returned."),
        )
    finally:
        os.unlink(audio_path)


@st.cache_resource
def local_whisper_model(model_name: str):
    from faster_whisper import WhisperModel

    return WhisperModel(model_name, device="cpu", compute_type="int8")


def process_audio_local(uploaded_file: Any, model_name: str, teacher_focus: str, target_language: str) -> dict[str, Any]:
    model = local_whisper_model(model_name)
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1] or ".wav") as audio_file:
        audio_file.write(uploaded_file.getvalue())
        audio_path = audio_file.name

    try:
        segments, _ = model.transcribe(audio_path)
        transcript = " ".join(segment.text.strip() for segment in segments).strip()
    finally:
        os.unlink(audio_path)

    prompt = (
        "Return JSON with exactly these keys: summary (string), action_items (array of objects "
        "with task, owner, due). Use 'Unassigned' or 'Not specified' when details are missing. "
        f"This teacher note focuses on {teacher_focus.lower()}. Write summary and tasks in {target_language}. "
        "Do not invent names, dates, or tasks.\n\n"
        f"Transcript:\n{transcript}"
    )
    request_body = json.dumps({
        "model": os.getenv("OLLAMA_MODEL", "llama3.2"),
        "stream": False,
        "format": "json",
        "messages": [
            {"role": "system", "content": "Extract concise, practical action items from teacher voice notes."},
            {"role": "user", "content": prompt},
        ],
    }).encode("utf-8")
    request = urllib.request.Request(
        "http://127.0.0.1:11434/api/chat",
        data=request_body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as error:
        raise RuntimeError("Ollama is not running. Install Ollama and start it before using local mode.") from error

    return parse_llm_result(payload.get("message", {}).get("content", "{}"), uploaded_file.name, transcript)


def pdf_bytes(results: list[dict[str, Any]]) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = io.BytesIO()
    document = SimpleDocTemplate(output, pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    story = [Paragraph("From Voice to Action", styles["Title"]), Spacer(1, 12)]
    for result in results:
        story.extend([Paragraph(result["file_name"], styles["Heading2"]), Paragraph(result["summary"], styles["BodyText"]), Spacer(1, 8)])
        story.append(Paragraph("Action items", styles["Heading3"]))
        for item in result["action_items"]:
            story.append(Paragraph(f"- {item.get('task', 'Untitled task')} | Owner: {item.get('owner', 'Unassigned')} | Due: {item.get('due', 'Not specified')}", styles["BodyText"]))
        story.extend([Spacer(1, 8), Paragraph("Raw transcript", styles["Heading3"]), Paragraph(result["transcript"], styles["BodyText"]), Spacer(1, 16)])
    document.build(story)
    return output.getvalue()


def text_export(results: list[dict[str, Any]]) -> str:
    lines = ["FROM VOICE TO ACTION", datetime.now().strftime("Generated %Y-%m-%d %H:%M"), ""]
    for result in results:
        lines.extend([result["file_name"].upper(), "", "EXECUTIVE SUMMARY", result["summary"], "", "ACTION ITEMS"])
        lines.extend([f"- {item.get('task', 'Untitled task')} | Owner: {item.get('owner', 'Unassigned')} | Due: {item.get('due', 'Not specified')}" for item in result["action_items"]])
        lines.extend(["", "RAW TRANSCRIPT", result["transcript"], "", "=" * 70, ""])
    return "\n".join(lines)


def word_bytes(results: list[dict[str, Any]]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("From Voice to Action", 0)
    for result in results:
        document.add_heading(result["file_name"], level=1)
        document.add_heading("Summary", level=2)
        document.add_paragraph(result["summary"])
        document.add_heading("Action items", level=2)
        for item in result["action_items"]:
            document.add_paragraph(
                f"{item.get('task', 'Untitled task')} | Owner: {item.get('owner', 'Unassigned')} | Due: {item.get('due', 'Not specified')}",
                style="List Bullet",
            )
        document.add_heading("Original transcript", level=2)
        document.add_paragraph(result["transcript"])
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def email_link(results: list[dict[str, Any]]) -> str:
    body = quote(text_export(results))
    return f"mailto:?subject=Teacher%20voice%20notes&body={body}"


def calendar_ics(results: list[dict[str, Any]]) -> str:
    today = datetime.now().date()
    events = []
    for result in results:
        for item_index, item in enumerate(result["action_items"]):
            due_text = str(item.get("due", "Not specified"))
            due_date = today
            due_lower = due_text.lower()
            weekdays = {name.lower(): index for index, name in enumerate(("Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"))}
            for weekday, weekday_index in weekdays.items():
                if weekday in due_lower:
                    due_date = today + timedelta(days=(weekday_index - today.weekday()) % 7 or 7)
                    break
            if "tomorrow" in due_lower:
                due_date = today + timedelta(days=1)
            elif "next week" in due_lower:
                due_date = today + timedelta(days=7)
            task = str(item.get("task", "Untitled task")).replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,")
            description = f"Owner: {item.get('owner', 'Unassigned')}\\nDue: {due_text}\\nSource: {result['file_name']}"
            events.append(
                "BEGIN:VEVENT\n"
                f"UID:{result['file_name']}-{item_index}@voice-notes\n"
                f"DTSTART;VALUE=DATE:{due_date.strftime('%Y%m%d')}\n"
                f"SUMMARY:{task}\n"
                f"DESCRIPTION:{description}\n"
                "END:VEVENT"
            )
    return "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Voice Notes//Teacher Tasks//EN\n" + "\n".join(events) + "\nEND:VCALENDAR\n"


st.set_page_config(page_title="From Voice to Action", page_icon="✦", layout="wide")
with st.sidebar:
    dark_mode = st.toggle("☾ / ☀", value=False, help="Switch between dark mode and warm white mode.")

theme = {
    "paper": "#101827" if dark_mode else "#fffdf7",
    "ink": "#f4f4f0" if dark_mode else "#111111",
    "muted": "#b8c4d4" if dark_mode else "#626262",
    "panel": "#24282c" if dark_mode else "#ffffff",
    "border": "#3b4247" if dark_mode else "#dce4dd",
    "task": "#173c3b" if dark_mode else "#edf8f5",
    "hero_ink": "#ffffff" if dark_mode else "#111111",
    "hero_border": "rgba(255,255,255,.3)" if dark_mode else "#d7dfd9",
}
st.markdown(
    f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=Space+Grotesk:wght@500;600;700&display=swap');
    :root {{ --ink: {theme['ink']}; --muted: {theme['muted']}; --cyan: #16b8c4; --coral: #ff735f; --paper: {theme['paper']}; --panel: {theme['panel']}; --border: {theme['border']}; --task: {theme['task']}; }}
    .stApp {{ background: var(--paper); color: var(--ink); font-family: 'DM Sans', sans-serif; }}
    h1, h2, h3, p, label, [data-testid='stCaptionContainer'] {{ font-family: 'Space Grotesk', sans-serif !important; color: var(--ink) !important; }}
    .hero h1, .hero p, .hero .eyebrow {{ color: {theme['hero_ink']} !important; }}
    h1 {{ font-size: clamp(2.2rem, 5vw, 4.6rem) !important; line-height: .98 !important; letter-spacing: -0.04em; }}
    .hero {{ padding: 2.5rem 0 1.5rem; border-bottom: 1px solid {theme['hero_border']}; }}
    .eyebrow {{ color: var(--coral); font-weight: 700; letter-spacing: .12em; text-transform: uppercase; font-size: .75rem; }}
    .subhead {{ color: var(--muted); font-size: 1.05rem; max-width: 720px; }}
    .panel {{ background: var(--panel); border: 1px solid var(--border); border-radius: 14px; padding: 1.2rem; box-shadow: 0 10px 30px rgba(20,43,53,.05); }}
    .result-label {{ color: var(--coral); font-size: .72rem; font-weight: 700; letter-spacing: .1em; text-transform: uppercase; }}
    .summary {{ font-family: 'Space Grotesk', sans-serif; font-size: 1.35rem; line-height: 1.25; }}
    .task {{ background: var(--task); border-left: 4px solid var(--cyan); padding: .75rem .9rem; margin: .55rem 0; border-radius: 0 8px 8px 0; }}
    .task strong {{ display: block; color: var(--ink); }}
    .meta {{ color: var(--muted); font-size: .84rem; }}
    [data-testid='stFileUploader'] {{ background: #f9faf6; border: 1px dashed #9eb8b4; border-radius: 10px; padding: .35rem; }}
    .stButton > button, .stDownloadButton > button {{ border-radius: 999px; font-weight: 700; }}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown('<div class="hero"><div class="eyebrow">From voice to action</div><h1>Speak it.<br>Act on it.</h1><p class="subhead">Upload or record a voice memo. Get the useful part back: a sharp summary, a clean transcript, and tasks you can actually act on.</p></div>', unsafe_allow_html=True)

if "user" not in st.session_state:
    st.session_state["user"] = None

if st.session_state["user"] is None:
    st.markdown("## Sign in to your classroom workspace")
    st.caption("Teachers manage notes and send messages. Students can read messages from their teacher.")
    with st.form("login_form"):
        login_role = st.selectbox("Login as", ["Teacher", "Student"])
        login_username = st.text_input("Username")
        login_password = st.text_input("Password", type="password")
        login_submitted = st.form_submit_button("Sign in", type="primary", use_container_width=True)
    if login_submitted:
        normalized_username = login_username.strip().lower()
        normalized_role = login_role.strip().lower()
        account = USERS.get(normalized_username)
        if account and account["password"] == login_password.strip() and account["role"].lower() == normalized_role:
            st.session_state["user"] = {"username": normalized_username, "role": account["role"]}
            st.rerun()
        st.error("The username, password, or role is incorrect.")
    st.info("Starter accounts: teacher1 / teacher123 and student1 / student123")
    st.stop()

user = st.session_state["user"]
messages = load_messages()
with st.sidebar:
    st.markdown(f"**{user['role']}:** {user['username']}")
    if st.button("Sign out", use_container_width=True):
        st.session_state["user"] = None
        st.rerun()

if user["role"] == "Teacher":
    st.markdown("### Teacher messages")
    with st.form("teacher_message_form", clear_on_submit=True):
        message_recipient = st.text_input("Send to student username", value="student1")
        message_subject = st.text_input("Message subject")
        message_body = st.text_area("Message", height=100)
        message_sent = st.form_submit_button("Send message", type="primary")
    if message_sent:
        if message_subject.strip() and message_body.strip():
            messages.append({
                "sender": user["username"],
                "recipient": message_recipient.strip() or "all",
                "subject": message_subject.strip(),
                "body": message_body.strip(),
                "created_at": datetime.now().isoformat(timespec="seconds"),
            })
            save_messages(messages)
            st.success("Message sent.")
        else:
            st.warning("Add a subject and message before sending.")
elif user["role"] == "Student":
    st.markdown("### Messages from your teacher")
    student_messages = [message for message in messages if message.get("recipient") in (user["username"], "all")]
    if student_messages:
        for message in reversed(student_messages):
            with st.expander(f"{message['subject']} · {message['created_at']}", expanded=True):
                st.write(message["body"])
                st.caption(f"From {message['sender']}")
    else:
        st.info("No teacher messages yet.")
    st.stop()

with st.sidebar:
    st.markdown("### Processing")
    processing_mode = st.radio("Processing mode", ["Demo / Test", "Free local", "Google Gemini cloud"], index=0)
    local_mode = processing_mode == "Free local"
    quality_options = {"Fast / lower accuracy": "tiny", "Balanced": "base", "Best / slower": "small"}
    whisper_quality = st.selectbox("Transcription quality", list(quality_options), index=1, disabled=not local_mode)
    teacher_focus = st.selectbox(
        "Teacher focus",
        ["General note", "Lesson planning", "Student meeting", "Parent communication", "Grading and feedback"],
        disabled=processing_mode == "Demo / Test",
    )
    target_language = st.selectbox(
        "Output language",
        ["English", "Hindi", "Spanish", "French", "German", "Tamil", "Telugu"],
        disabled=processing_mode == "Demo / Test",
    )
    if processing_mode == "Google Gemini cloud":
        if server_api_key("GEMINI_API_KEY"):
            st.caption("Gemini processes audio on the cloud for every device.")
        else:
            st.warning("Add GEMINI_API_KEY in Streamlit App Settings > Secrets before using cloud mode.")
    elif local_mode:
        st.caption("Free local mode needs Ollama and the llama3.2 model on this computer.")
    else:
        st.caption("Demo mode lets your teacher test the complete interface without payment or setup.")

st.markdown("### 01 / Bring in a note")
recorded_audio = st.audio_input("Record a voice note", sample_rate=44100)
uploaded_files = st.file_uploader("Or drop audio files here", type=["mp3", "wav", "m4a"], accept_multiple_files=True)
with st.expander("Recording tips for clearer text"):
    st.markdown("- Speak close to the microphone and use short, complete sentences.\n- Record in a quiet room and pause briefly between topics.\n- Say names, dates, and assignments slowly and clearly.\n- For a lesson note, mention the class, topic, and next step.")

audio_inputs = []
if recorded_audio is not None:
    audio_inputs.append(recorded_audio)
audio_inputs.extend(uploaded_files or [])

if audio_inputs:
    st.caption(f"{len(audio_inputs)} note{'s' if len(audio_inputs) != 1 else ''} ready to process")
    if st.button("Transcribe and extract actions", type="primary", use_container_width=True):
        results = []
        progress = st.progress(0, text="Starting...")
        for index, audio_file in enumerate(audio_inputs):
            try:
                if processing_mode == "Demo / Test":
                    result = demo_result(getattr(audio_file, "name", "recorded-note.wav"))
                elif processing_mode == "Google Gemini cloud":
                    result = process_audio_gemini(audio_file, teacher_focus, target_language)
                else:
                    result = process_audio_local(audio_file, quality_options[whisper_quality], teacher_focus, target_language)
                result["created_at"] = datetime.now().isoformat(timespec="seconds")
                results.append(result)
                progress.progress((index + 1) / len(audio_inputs), text=f"Processed {index + 1} of {len(audio_inputs)}")
            except Exception as error:
                st.error(f"Could not process {getattr(audio_file, 'name', 'audio file')}: {error}")
        st.session_state["results"] = results
        progress.empty()
else:
    st.info("Add an audio file or record a note to begin.")

if "history" not in st.session_state:
    st.session_state["history"] = load_history()
if results := st.session_state.get("results", []):
    for result in results:
        if result not in st.session_state["history"]:
            st.session_state["history"].append(result)
    save_history(st.session_state["history"])

st.markdown("### Note history")
history_query = st.text_input("Search by student, class, subject, date, or keyword", placeholder="Try: maths, Priya, Friday")
history = st.session_state["history"]
if history_query:
    query = history_query.lower()
    history = [
        result for result in history
        if query in json.dumps(result, ensure_ascii=False).lower()
    ]
results = history
if results:
    st.markdown("### 02 / Your organized notes")
    for result in results:
        st.markdown(f"#### {result['file_name']}")
        delete_key = f"delete_{result['file_name']}"
        if st.button("Delete note", key=delete_key, help="Remove this saved voice note from local history."):
            st.session_state["history"] = [
                saved_result for saved_result in st.session_state["history"]
                if saved_result.get("file_name") != result.get("file_name")
            ]
            save_history(st.session_state["history"])
            st.session_state["results"] = []
            st.rerun()
        summary_col, actions_col = st.columns([1, 1], gap="large")
        with summary_col:
            st.markdown('<div class="panel"><div class="result-label">Executive summary</div><p class="summary">' + result["summary"] + '</p></div>', unsafe_allow_html=True)
            with st.expander("View raw transcript", expanded=True):
                st.write(result["transcript"])
        with actions_col:
            st.markdown('<div class="panel"><div class="result-label">Action items</div>', unsafe_allow_html=True)
            if result["action_items"]:
                for item_index, item in enumerate(result["action_items"]):
                    task_key = f"completed_{result['file_name']}_{item_index}"
                    if task_key not in st.session_state:
                        st.session_state[task_key] = item.get("completed", False)
                    completed = st.checkbox(item.get("task", "Untitled task"), key=task_key)
                    item["completed"] = completed
                    for saved_result in st.session_state["history"]:
                        if saved_result.get("file_name") == result.get("file_name"):
                            saved_result["action_items"] = result["action_items"]
                    save_history(st.session_state["history"])
                    status = "Completed" if completed else "Open"
                    st.markdown(
                        f'<div class="task"><span class="meta">{item.get("owner", "Unassigned")} · due {item.get("due", "Not specified")} · {status}</span></div>',
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("No action items found.")
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("### 03 / Take it with you")
    export_col_1, export_col_2, export_col_3 = st.columns(3)
    with export_col_1:
        st.download_button("Download .txt note", text_export(results), file_name="voice-notes-action-items.txt", mime="text/plain", use_container_width=True)
    with export_col_2:
        st.download_button("Download PDF", pdf_bytes(results), file_name="voice-notes-action-items.pdf", mime="application/pdf", use_container_width=True)
    with export_col_3:
        st.download_button("Download Word", word_bytes(results), file_name="teacher-notes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    st.download_button("Download calendar deadlines", calendar_ics(results), file_name="teacher-deadlines.ics", mime="text/calendar", use_container_width=True)
    st.markdown("Download the Word file and upload it to Google Docs to edit and share it.")
    st.markdown(f"[Open email draft](<{email_link(results)}>)")